import os
from docx import Document
import win32com.client as win32

# Create a Word document using python-docx
doc = Document()
doc.add_heading('Document Title', 0)
doc.add_paragraph('A paragraph in the document.')
doc_path = 'example.docm'  # Save as .docm to enable macros
doc.save(doc_path)

# Define the path to the macro file (a text file containing VBA code)
macro_file_path = 'macro.txt'

# Read the macro content from the text file
with open(macro_file_path, 'r') as file:
    macro_code = file.read()

# Add the macro to the Word document using win32com.client
word = win32.Dispatch('Word.Application')
word.Visible = False

# Open the document
doc = word.Documents.Open(os.path.abspath(doc_path))

# Add the macro to the document
word_module = doc.VBProject.VBComponents.Add(1)  # 1 indicates a standard module
word_module.CodeModule.AddFromString(macro_code)

# Save and close the document
doc.Save()
doc.Close()

# Quit the Word application
word.Quit()
